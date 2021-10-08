from docx import Document
from docx.shared import Inches
import datetime
import randomee

x = datetime.datetime.now()
y = datetime.datetime.now()
d = y.strftime("%d")
dada = int(d)

def week(dada):
    max = 31
    min = 0
    if dada > max:
        dada = d +7
    else:
        dada = min
        dada = dada +7
    return dada
dd = week(dada)
#print(dd)

datum = x.strftime(f"von %d-%m-%Y bis {dd}-%m-%Y")
heute = x.strftime(f'{dd}-%m-%Y')
nachweiss_nummer = 0
woche_nummer = 0
document = Document()
document.add_paragraph(f'Name: Albyama        Ausbildungnachweis Nr: {nachweiss_nummer}            {datum}     ')

p = document.add_paragraph(f'Firma: \t \t Ausbildung Jahr 2 \t \t Woche: {woche_nummer}\nStrasse')
p = document.add_paragraph('Ausbildung Abteilung:   ')



document.add_heading('Betribliche Tätigkeiten.', level=1)
#document.add_paragraph('Intense quote', style='Intense Quote')

y = randomee.aggiungi_a_beri(6)
for i in range(1):
    for x in y[0]:
        document.add_paragraph(
            f'{x}', style='List Bullet')

document.add_heading('Beschreibung eines Arbeitsvorganges dieser Woche:', level=1)
#qua arriva la parte dove json file oder story

for i in range(1):
    sscelta = randomee.tema_beri(y[1])
        #print(f'eeeccoo {x}')
    document.add_paragraph(
        f'{sscelta}', style='List Bullet')


document.add_heading('Berufsschule (Themen des Unterrichts)')

for i in range(1):
    for x in randomee.aggiungi_a_beri(5):
        document.add_paragraph(
            f'{x}', style='List Bullet')


p = document.add_paragraph(f'\n\n\n Für die Richtigkeit\n\nDatum {heute} \t\t\t\t\t\tDatum {heute}')
p = document.add_paragraph(f'Auszubildender\t\t\t\t\t\tAusbilder')
document.add_page_break()

document.save('demo.docx')